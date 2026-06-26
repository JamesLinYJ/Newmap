# +-------------------------------------------------------------------------
#
#   地理智能平台 - 气象 DOCX 报告生成
#
#   文件:       report.py
#
#   日期:       2026年06月25日
#   作者:       OpenAI Codex
# --------------------------------------------------------------------------

# 模块职责
#
# 将已经校验过的气象 metadata、统计事实和模型解读引用写入 DOCX。
# 这里不调用模型、不补造分析结论；缺少显式解读时由上游服务直接失败。

from __future__ import annotations

from pathlib import Path
from typing import Any


def write_meteorological_report_docx(
    *,
    output_path: Path,
    dataset_id: str | None,
    filename: str,
    metadata: dict[str, Any],
    stats_rows: list[dict[str, Any]],
    llm_interpretation: str,
    generated_at: str,
) -> dict[str, Any]:
    """Write a compact DOCX report from explicit meteorological facts."""

    if not llm_interpretation.strip():
        raise ValueError("气象报告缺少经过校验的模型解读引用。")
    docx = _docx()
    document = docx.Document()
    document.add_heading("气象数据分析报告", level=0)
    document.add_paragraph(f"生成时间：{generated_at}")
    document.add_paragraph(f"数据文件：{filename}")
    if dataset_id:
        document.add_paragraph(f"数据集 ID：{dataset_id}")

    document.add_heading("一、数据概况", level=1)
    overview = document.add_table(rows=1, cols=2)
    overview.style = "Table Grid"
    overview.rows[0].cells[0].text = "项目"
    overview.rows[0].cells[1].text = "内容"
    _append_row(overview, "格式", str(metadata.get("format") or metadata.get("engine") or "未知"))
    _append_row(overview, "地理范围", _format_bounds(metadata.get("bounds")))
    _append_row(overview, "变量数量", str(len(_variables(metadata))))
    warnings = metadata.get("warnings")
    if isinstance(warnings, list) and warnings:
        _append_row(overview, "注意事项", "；".join(str(item) for item in warnings[:5]))

    document.add_heading("二、变量清单", level=1)
    variables = _variables(metadata)
    if variables:
        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ["变量", "单位", "维度", "值域", "地图能力"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for variable in variables:
            _append_variable_row(table, variable)
    else:
        document.add_paragraph("metadata 中未包含可展示变量。")

    document.add_heading("三、统计摘要", level=1)
    usable_stats = [row for row in stats_rows if isinstance(row, dict)]
    if usable_stats:
        table = document.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        headers = ["变量", "最小值", "最大值", "平均值", "P50", "P90", "说明"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for row in usable_stats:
            cells = table.add_row().cells
            cells[0].text = str(row.get("variable") or "")
            cells[1].text = _number(row.get("min"))
            cells[2].text = _number(row.get("max"))
            cells[3].text = _number(row.get("mean"))
            cells[4].text = _number(row.get("p50") or row.get("median"))
            cells[5].text = _number(row.get("p90"))
            cells[6].text = str(row.get("error") or row.get("unit") or "")
    else:
        document.add_paragraph("未生成变量统计摘要。")

    document.add_heading("四、模型解读", level=1)
    for paragraph in [item.strip() for item in llm_interpretation.splitlines() if item.strip()]:
        document.add_paragraph(paragraph)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return {
        "outputPath": str(output_path),
        "filename": filename,
        "datasetId": dataset_id,
        "generatedAt": generated_at,
        "variableCount": len(variables),
        "statsRowCount": len(usable_stats),
    }


def _append_row(table: Any, key: str, value: str) -> None:
    cells = table.add_row().cells
    cells[0].text = key
    cells[1].text = value


def _append_variable_row(table: Any, variable: dict[str, Any]) -> None:
    cells = table.add_row().cells
    cells[0].text = str(variable.get("name") or "")
    cells[1].text = str(variable.get("unit") or "")
    dimensions = variable.get("dimensions")
    cells[2].text = ", ".join(str(item) for item in dimensions) if isinstance(dimensions, list) else ""
    cells[3].text = _format_range(variable.get("valueRange"))
    cells[4].text = "可制图" if variable.get("mapReady") else "不可制图"


def _variables(metadata: dict[str, Any]) -> list[dict[str, Any]]:
    raw = metadata.get("variables")
    if not isinstance(raw, list):
        return []
    return [item for item in raw if isinstance(item, dict)]


def _format_bounds(value: Any) -> str:
    if not isinstance(value, list) or len(value) != 4:
        return "未提供"
    return ", ".join(_number(item) for item in value)


def _format_range(value: Any) -> str:
    if not isinstance(value, list) or len(value) != 2:
        return ""
    return f"{_number(value[0])} ~ {_number(value[1])}"


def _number(value: Any) -> str:
    try:
        return f"{float(value):.4g}"
    except (TypeError, ValueError):
        return ""


def _docx() -> Any:
    try:
        import docx

        return docx
    except Exception as exc:  # noqa: BLE001 - dependency import failure must surface clearly.
        raise RuntimeError("python-docx 不可用，无法生成气象 DOCX 报告。") from exc
