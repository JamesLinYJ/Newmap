# +-------------------------------------------------------------------------
#
#   地理智能平台 - 气象 DOCX 报告生成
#
#   文件:       report.py
#
#   日期:       2026年05月26日
#   作者:       OpenAI Codex
# --------------------------------------------------------------------------

# 模块职责
#
# 将气象数据 inspect metadata 和统计摘要整理成正式 DOCX 解读报告。
# 报告生成只消费已经解析出的事实，不回扫运行历史，也不伪造分析结论。

from __future__ import annotations

from pathlib import Path
from typing import Any


def write_weather_report_docx(
    *,
    output_path: Path,
    dataset_id: str | None,
    filename: str,
    metadata: dict[str, Any],
    stats_rows: list[dict[str, Any]],
    llm_interpretation: str,
    generated_at: str,
) -> dict[str, Any]:
    # DOCX 排版入口。
    #
    # python-docx 写出的文档保持 Word/WPS 兼容；样式统一在这里声明，
    # 避免 API 或工具层散落格式细节。
    docx = _docx()
    document = docx.Document()
    _configure_document(document)

    title = "NC 气象数据解读报告" if _is_netcdf_metadata(metadata, filename) else "气象数据解读报告"
    _add_title(document, title, filename)
    _add_key_value_table(
        document,
        "一、数据概览",
        [
            ("文件名", filename),
            ("数据集 ID", dataset_id or "未记录"),
            ("格式 / 引擎", f"{metadata.get('format') or 'unknown'} / {metadata.get('engine') or 'unknown'}"),
            ("生成时间", generated_at),
            ("空间范围", _format_bounds(metadata.get("bounds"))),
            ("坐标字段", _format_coordinates(metadata.get("coordinates"))),
        ],
    )

    variables = _metadata_variables(metadata)
    _add_summary(document, metadata, variables)
    _add_variable_table(document, variables)
    _add_dimension_section(document, metadata, variables)
    _add_stats_table(document, stats_rows)
    _add_llm_interpretation(document, llm_interpretation)
    _add_backend_section(document, metadata, variables)
    _add_warning_section(document, metadata, stats_rows)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return {
        "title": title,
        "filename": filename,
        "datasetId": dataset_id,
        "variableCount": len(variables),
        "statsCount": len([item for item in stats_rows if not item.get("error")]),
        "warningCount": len(_metadata_warnings(metadata)) + len([item for item in stats_rows if item.get("error")]),
        "format": metadata.get("format"),
        "engine": metadata.get("engine"),
        "llmInterpretationChars": len(llm_interpretation),
    }


def _configure_document(document: Any) -> None:
    styles = document.styles
    normal = styles["Normal"]
    _set_run_font(normal.font, size=10.5)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = _pt(4)

    for style_name, size, bold, color in (
        ("Title", 22, True, "1F4E79"),
        ("Heading 1", 14, True, "1F4E79"),
        ("Heading 2", 12, True, "2F5597"),
    ):
        style = styles[style_name]
        _set_run_font(style.font, size=size, bold=bold, color=color)
        style.paragraph_format.space_before = _pt(10)
        style.paragraph_format.space_after = _pt(6)

    section = document.sections[0]
    section.top_margin = _cm(1.8)
    section.bottom_margin = _cm(1.6)
    section.left_margin = _cm(1.9)
    section.right_margin = _cm(1.9)


def _add_title(document: Any, title: str, filename: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = _wd_align().CENTER
    run = paragraph.add_run(title)
    _set_run_font(run.font, size=22, bold=True, color="1F4E79")
    subtitle = document.add_paragraph()
    subtitle.alignment = _wd_align().CENTER
    subtitle_run = subtitle.add_run(filename)
    _set_run_font(subtitle_run.font, size=11, color="666666")


def _add_summary(document: Any, metadata: dict[str, Any], variables: list[dict[str, Any]]) -> None:
    document.add_heading("二、自动解读摘要", level=1)
    variable_count = len(variables)
    analysis_ready = sum(1 for item in variables if item.get("analysisReady"))
    map_ready = sum(1 for item in variables if item.get("mapReady"))
    preferred = _preferred_backend_counts(variables)
    text = (
        f"本报告识别到 {variable_count} 个数值型气象变量，其中 {analysis_ready} 个可进行统计分析，"
        f"{map_ready} 个具备地图栅格展示条件。"
    )
    if preferred:
        text += " 后端路由建议：" + "，".join(f"{name} {count} 个变量" for name, count in preferred.items()) + "。"
    if metadata.get("bounds"):
        text += f" 数据空间范围为 {_format_bounds(metadata.get('bounds'))}。"
    _add_body_paragraph(document, text)


def _add_variable_table(document: Any, variables: list[dict[str, Any]]) -> None:
    document.add_heading("三、变量清单", level=1)
    table = _create_table(document, ["变量", "含义", "维度", "形状", "单位", "能力", "推荐后端"])
    for variable in variables:
        capability = []
        if variable.get("analysisReady"):
            capability.append("可统计")
        if variable.get("mapReady"):
            capability.append("可制图")
        _add_table_row(
            table,
            [
                str(variable.get("name") or "-"),
                str(variable.get("longName") or "-"),
                " × ".join(str(item) for item in variable.get("dimensions") or []) or "-",
                " × ".join(str(item) for item in variable.get("shape") or []) or "-",
                str(variable.get("unit") or "-"),
                " / ".join(capability) or "仅元数据",
                str(variable.get("preferredBackend") or "-"),
            ],
        )


def _add_dimension_section(document: Any, metadata: dict[str, Any], variables: list[dict[str, Any]]) -> None:
    document.add_heading("四、时间与垂直层", level=1)
    times = metadata.get("times") if isinstance(metadata.get("times"), list) else []
    levels = metadata.get("levels") if isinstance(metadata.get("levels"), list) else []
    max_time_count = max((int(item.get("timeCount") or 0) for item in variables), default=0)
    max_level_count = max((int(item.get("levelCount") or 0) for item in variables), default=0)
    _add_body_paragraph(document, f"时间片数量：{max_time_count}；垂直层数量：{max_level_count}。")
    if times:
        _add_body_paragraph(document, "时间片示例：" + "、".join(str(item) for item in times[:8]))
    if levels:
        _add_body_paragraph(document, "垂直层示例：" + "、".join(str(item) for item in levels[:8]))


def _add_stats_table(document: Any, stats_rows: list[dict[str, Any]]) -> None:
    document.add_heading("五、统计摘要", level=1)
    if not stats_rows:
        _add_body_paragraph(document, "当前报告未生成统计摘要。")
        return
    table = _create_table(document, ["变量", "有效值", "最小", "最大", "平均", "中位数", "P90", "备注"])
    for item in stats_rows:
        if item.get("error"):
            _add_table_row(table, [str(item.get("variable") or "-"), "-", "-", "-", "-", "-", "-", str(item["error"])])
            continue
        _add_table_row(
            table,
            [
                str(item.get("variable") or "-"),
                str(item.get("count") or 0),
                _format_number(item.get("min"), item.get("unit")),
                _format_number(item.get("max"), item.get("unit")),
                _format_number(item.get("mean"), item.get("unit")),
                _format_number(item.get("median", item.get("p50")), item.get("unit")),
                _format_number(item.get("p90"), item.get("unit")),
                _format_stat_context(item),
            ],
        )


def _add_backend_section(document: Any, metadata: dict[str, Any], variables: list[dict[str, Any]]) -> None:
    document.add_heading("七、地图执行与后端能力", level=1)
    backend_summary = metadata.get("backendSummary") if isinstance(metadata.get("backendSummary"), dict) else {}
    if backend_summary:
        _add_body_paragraph(document, "后端元数据：" + _compact_json(backend_summary))
    table = _create_table(document, ["变量", "xarray 分析", "rasterio/GDAL 制图", "bounds"])
    for variable in variables:
        backends = variable.get("backends") if isinstance(variable.get("backends"), list) else []
        raster_backend = next((item for item in backends if isinstance(item, dict) and item.get("name") == "rasterio"), None)
        _add_table_row(
            table,
            [
                str(variable.get("name") or "-"),
                "是" if variable.get("analysisReady") else "否",
                "是" if raster_backend and raster_backend.get("mapReady") else "否",
                _format_bounds(variable.get("bounds")),
            ],
        )


def _add_warning_section(document: Any, metadata: dict[str, Any], stats_rows: list[dict[str, Any]]) -> None:
    warnings = _metadata_warnings(metadata)
    warnings.extend(f"{item.get('variable')}: {item.get('error')}" for item in stats_rows if item.get("error"))
    document.add_heading("八、限制与注意事项", level=1)
    if not warnings:
        _add_body_paragraph(document, "未发现阻断性警告。后续制图和统计仍应以用户指定变量、time、level 与 bbox 为准。")
        return
    for warning in warnings:
        _add_body_paragraph(document, f"• {warning}")


def _add_llm_interpretation(document: Any, text: str) -> None:
    document.add_heading("六、大模型综合解读", level=1)
    for paragraph_text in _split_interpretation(text):
        _add_body_paragraph(document, paragraph_text)


def _add_key_value_table(document: Any, heading: str, rows: list[tuple[str, str]]) -> None:
    document.add_heading(heading, level=1)
    table = _create_table(document, ["项目", "内容"])
    for key, value in rows:
        _add_table_row(table, [key, value])


def _create_table(document: Any, headers: list[str]) -> Any:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = _wd_table_alignment().CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        cell = header_cells[index]
        _shade_cell(cell, "D9EAF7")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = _wd_align().CENTER
        run = paragraph.add_run(header)
        _set_run_font(run.font, size=9.5, bold=True, color="1F4E79")
    return table


def _add_table_row(table: Any, values: list[str]) -> None:
    row = table.add_row()
    for index, value in enumerate(values):
        cell = row.cells[index]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = _wd_align().CENTER if index != 1 else _wd_align().LEFT
        run = paragraph.add_run(value)
        _set_run_font(run.font, size=9)


def _add_body_paragraph(document: Any, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = _cm(0.74) if not text.startswith("•") else None
    run = paragraph.add_run(text)
    _set_run_font(run.font, size=10.5)


def _set_run_font(font: Any, *, size: float, bold: bool | None = None, color: str | None = None) -> None:
    font.name = "Microsoft YaHei"
    font.size = _pt(size)
    if bold is not None:
        font.bold = bold
    if color:
        font.color.rgb = _rgb(color)
    element = getattr(font, "element", None)
    if element is not None and getattr(element, "rPr", None) is not None:
        element.rPr.rFonts.set(_qn("w:eastAsia"), "Microsoft YaHei")


def _shade_cell(cell: Any, fill: str) -> None:
    shading = _oxml_element("w:shd")
    shading.set(_qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def _metadata_variables(metadata: dict[str, Any]) -> list[dict[str, Any]]:
    raw_variables = metadata.get("variables")
    if not isinstance(raw_variables, list):
        return []
    return [item for item in raw_variables if isinstance(item, dict)]


def _metadata_warnings(metadata: dict[str, Any]) -> list[str]:
    warnings = metadata.get("warnings")
    if not isinstance(warnings, list):
        return []
    return [str(item) for item in warnings if str(item).strip()]


def _split_interpretation(text: str) -> list[str]:
    paragraphs = [item.strip() for item in text.replace("\r\n", "\n").split("\n") if item.strip()]
    return paragraphs or ["大模型未提供有效解读。"]


def _preferred_backend_counts(variables: list[dict[str, Any]]) -> dict[str, int]:
    counts: dict[str, int] = {}
    for variable in variables:
        backend = str(variable.get("preferredBackend") or "").strip()
        if backend and backend != "none":
            counts[backend] = counts.get(backend, 0) + 1
    return counts


def _format_coordinates(value: Any) -> str:
    if not isinstance(value, dict):
        return "未记录"
    return "；".join(f"{key}={item or '-'}" for key, item in value.items())


def _format_bounds(value: Any) -> str:
    if isinstance(value, (list, tuple)) and len(value) >= 4:
        try:
            return ", ".join(f"{float(item):.4f}" for item in value[:4])
        except (TypeError, ValueError):
            return str(value)
    return "无地图范围"


def _format_number(value: Any, unit: Any = None) -> str:
    if value is None:
        return "-"
    try:
        text = f"{float(value):.3f}".rstrip("0").rstrip(".")
    except (TypeError, ValueError):
        text = str(value)
    unit_text = str(unit).strip() if unit else ""
    return f"{text} {unit_text}" if unit_text else text


def _format_stat_context(item: dict[str, Any]) -> str:
    parts: list[str] = []
    if item.get("timeValue"):
        parts.append(f"time={item['timeValue']}")
    if item.get("levelValue"):
        parts.append(f"level={item['levelValue']}")
    return "；".join(parts) if parts else "默认切片"


def _compact_json(value: Any) -> str:
    import json

    return json.dumps(value, ensure_ascii=False, separators=(",", ":"))[:700]


def _is_netcdf_metadata(metadata: dict[str, Any], filename: str) -> bool:
    suffix = Path(filename).suffix.casefold()
    return suffix in {".nc", ".nc4"} or str(metadata.get("format") or "").casefold() == "netcdf"


def _docx() -> Any:
    import docx

    return docx


def _pt(value: float) -> Any:
    from docx.shared import Pt

    return Pt(value)


def _cm(value: float) -> Any:
    from docx.shared import Cm

    return Cm(value)


def _rgb(value: str) -> Any:
    from docx.shared import RGBColor

    value = value.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def _wd_align() -> Any:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    return WD_ALIGN_PARAGRAPH


def _wd_table_alignment() -> Any:
    from docx.enum.table import WD_TABLE_ALIGNMENT

    return WD_TABLE_ALIGNMENT


def _oxml_element(name: str) -> Any:
    from docx.oxml import OxmlElement

    return OxmlElement(name)


def _qn(name: str) -> str:
    from docx.oxml.ns import qn

    return qn(name)
